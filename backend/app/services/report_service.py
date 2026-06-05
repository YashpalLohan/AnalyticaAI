"""
AnalyticaAI — Report Service
Generates professional PDF and DOCX reports from dataset insights.
Uses reportlab for PDF and python-docx for Word documents.
"""
from __future__ import annotations

import io
import logging
from datetime import date
from typing import Any

from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.dataset import Dataset
from app.models.dataset_profile import DatasetProfile
from app.services.dataset_service import get_dataset
from app.services.insight_service import generate_insights, get_insights

logger = logging.getLogger(__name__)

# ── Severity colours (for PDF accent bars) ────────────────────────────────
SEVERITY_COLORS = {
    "high":   (217, 64,  64),   # error red
    "medium": (196, 122, 0),    # warning amber
    "low":    (30,  138, 82),   # success green
}
TYPE_COLORS = {
    "trend":          (79,  110, 247),  # blue
    "risk":           (217, 64,  64),   # red
    "opportunity":    (30,  138, 82),   # green
    "recommendation": (124, 58,  237),  # purple
}

# ── PDF generation ─────────────────────────────────────────────────────────

def _generate_pdf(
    dataset_name: str,
    row_count: int,
    col_count: int,
    health_score: float | None,
    insights: dict,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    W = A4[0] - 40 * mm  # usable width

    # ── Custom styles ──
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=22,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#16191F"),
        spaceAfter=6,
        leading=26,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#6B7080"),
        spaceAfter=4,
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontSize=13,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#16191F"),
        spaceBefore=14,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#363A45"),
        leading=15,
        spaceAfter=8,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontSize=8,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#6B7080"),
        spaceAfter=2,
    )
    insight_title_style = ParagraphStyle(
        "InsightTitle",
        parent=styles["Normal"],
        fontSize=11,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#16191F"),
        spaceAfter=3,
    )
    insight_body_style = ParagraphStyle(
        "InsightBody",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#363A45"),
        leading=14,
    )

    story = []

    # ── Cover ──
    story.append(Spacer(1, 10 * mm))
    story.append(Paragraph("Analytics Report", label_style))
    story.append(Paragraph(dataset_name, title_style))
    story.append(Paragraph(f"Generated on {date.today().strftime('%B %d, %Y')}", subtitle_style))
    story.append(HRFlowable(width=W, thickness=2, color=colors.HexColor("#4F6EF7"), spaceAfter=12))

    # ── Executive summary ──
    summary = insights.get("executive_summary", "")
    if summary:
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(summary, body_style))

    # ── Dataset overview table ──
    story.append(Paragraph("Dataset Overview", heading_style))
    overview_data = [
        ["Metric", "Value"],
        ["Dataset Name",    dataset_name],
        ["Total Rows",      f"{row_count:,}"],
        ["Total Columns",   str(col_count)],
        ["Health Score",    f"{health_score}/100" if health_score else "Not profiled"],
        ["Report Date",     date.today().strftime("%Y-%m-%d")],
    ]
    overview_table = Table(overview_data, colWidths=[W * 0.4, W * 0.6])
    overview_table.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0),  colors.HexColor("#16191F")),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F0EEE9"), colors.white]),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.HexColor("#DEDAD2")),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING",   (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 6),
    ]))
    story.append(overview_table)

    # ── Insights ──
    insight_list = insights.get("insights", [])
    if insight_list:
        story.append(Paragraph("Key Insights", heading_style))

        for ins in insight_list:
            ins_type = ins.get("type", "trend").lower()
            severity = ins.get("severity", "medium").lower()
            type_rgb  = TYPE_COLORS.get(ins_type, (79, 110, 247))
            type_hex  = "#{:02X}{:02X}{:02X}".format(*type_rgb)

            # Type badge + title in a mini-table for the accent bar effect
            badge_text = f"<font color='{type_hex}'><b>{ins_type.upper()}</b></font>"
            badge_para = Paragraph(badge_text, label_style)
            title_para = Paragraph(ins.get("title", ""), insight_title_style)
            desc_para  = Paragraph(ins.get("description", ""), insight_body_style)

            sev_rgb = SEVERITY_COLORS.get(severity, (196, 122, 0))
            sev_hex = "#{:02X}{:02X}{:02X}".format(*sev_rgb)
            sev_text = f"<font color='{sev_hex}'><b>{severity.upper()}</b></font>"
            sev_para = Paragraph(sev_text, label_style)

            card_data = [[badge_para, sev_para], [title_para, ""], [desc_para, ""]]
            card_table = Table(card_data, colWidths=[W * 0.82, W * 0.18])
            card_table.setStyle(TableStyle([
                ("BACKGROUND",  (0, 0), (-1, -1), colors.HexColor("#F0EEE9")),
                ("BOX",         (0, 0), (-1, -1), 1,   colors.HexColor("#DEDAD2")),
                ("LINEBEFORE",  (0, 0), (0, -1),  3,   colors.HexColor(type_hex)),
                ("LEFTPADDING",  (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING",   (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 6),
                ("SPAN",        (0, 1), (1, 1)),
                ("SPAN",        (0, 2), (1, 2)),
                ("VALIGN",      (0, 0), (-1, -1), "TOP"),
                ("ALIGN",       (1, 0), (1, 0),   "RIGHT"),
            ]))
            story.append(card_table)
            story.append(Spacer(1, 4 * mm))

    doc.build(story)
    return buffer.getvalue()


# ── DOCX generation ────────────────────────────────────────────────────────

def _generate_docx(
    dataset_name: str,
    row_count: int,
    col_count: int,
    health_score: float | None,
    insights: dict,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    TYPE_RGB_DOCX = {
        "trend":          RGBColor(79,  110, 247),
        "risk":           RGBColor(217, 64,  64),
        "opportunity":    RGBColor(30,  138, 82),
        "recommendation": RGBColor(124, 58,  237),
    }

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Cover ──
    label = doc.add_paragraph("ANALYTICS REPORT")
    label.runs[0].font.size  = Pt(9)
    label.runs[0].font.color.rgb = RGBColor(107, 112, 128)
    label.runs[0].font.bold = True
    label.paragraph_format.space_after = Pt(4)

    title = doc.add_heading(dataset_name, level=1)
    title.runs[0].font.size  = Pt(22)
    title.runs[0].font.color.rgb = RGBColor(22, 25, 31)
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph(f"Generated on {date.today().strftime('%B %d, %Y')}")
    sub.runs[0].font.size  = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(107, 112, 128)

    doc.add_paragraph()

    # ── Executive Summary ──
    summary = insights.get("executive_summary", "")
    if summary:
        h = doc.add_heading("Executive Summary", level=2)
        h.runs[0].font.color.rgb = RGBColor(22, 25, 31)
        p = doc.add_paragraph(summary)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(54, 58, 69)

    # ── Dataset Overview table ──
    h2 = doc.add_heading("Dataset Overview", level=2)
    h2.runs[0].font.color.rgb = RGBColor(22, 25, 31)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Dataset Name",  dataset_name),
        ("Total Rows",    f"{row_count:,}"),
        ("Total Columns", str(col_count)),
        ("Health Score",  f"{health_score}/100" if health_score else "Not profiled"),
        ("Report Date",   date.today().strftime("%Y-%m-%d")),
    ]
    for i, (label_txt, val_txt) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label_txt
        row.cells[1].text = val_txt
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Insights ──
    insight_list = insights.get("insights", [])
    if insight_list:
        h3 = doc.add_heading("Key Insights", level=2)
        h3.runs[0].font.color.rgb = RGBColor(22, 25, 31)

        for ins in insight_list:
            ins_type = ins.get("type", "trend").lower()
            severity = ins.get("severity", "medium").lower()
            type_color = TYPE_RGB_DOCX.get(ins_type, RGBColor(79, 110, 247))

            # Type badge
            badge = doc.add_paragraph()
            badge_run = badge.add_run(f"{ins_type.upper()}  ·  {severity.upper()}")
            badge_run.font.size  = Pt(8)
            badge_run.font.bold  = True
            badge_run.font.color.rgb = type_color
            badge.paragraph_format.space_before = Pt(8)
            badge.paragraph_format.space_after  = Pt(2)

            # Title
            t = doc.add_paragraph()
            t_run = t.add_run(ins.get("title", ""))
            t_run.font.size  = Pt(11)
            t_run.font.bold  = True
            t_run.font.color.rgb = RGBColor(22, 25, 31)
            t.paragraph_format.space_after = Pt(2)

            # Description
            d = doc.add_paragraph(ins.get("description", ""))
            d.runs[0].font.size  = Pt(10)
            d.runs[0].font.color.rgb = RGBColor(54, 58, 69)
            d.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public service functions ───────────────────────────────────────────────

async def generate_report(
    dataset_id: str,
    user_id: str,
    fmt: str,
    db: AsyncSession,
) -> tuple[bytes, str]:
    """
    Generate a PDF or DOCX report for the dataset.
    Returns (file_bytes, filename).
    fmt: 'pdf' | 'docx'
    """
    dataset: Dataset = await get_dataset(dataset_id, user_id, db)

    # Get insights (generate if not cached)
    try:
        insights = await get_insights(dataset_id, user_id, db)
    except HTTPException:
        insights = await generate_insights(dataset_id, user_id, db)

    # Get profile metadata
    row_count  = dataset.total_rows or 0
    col_count  = dataset.total_columns or 0
    health_score = None

    profile_result = await db.execute(
        select(DatasetProfile).where(DatasetProfile.dataset_id == dataset_id)
    )
    profile_orm = profile_result.scalar_one_or_none()
    if profile_orm:
        health_score = profile_orm.health_score

    safe_name = dataset.name.replace(" ", "_")[:40]

    try:
        if fmt == "pdf":
            content = _generate_pdf(
                dataset.name, row_count, col_count, health_score, insights
            )
            filename = f"{safe_name}_report.pdf"
        elif fmt == "docx":
            content = _generate_docx(
                dataset.name, row_count, col_count, health_score, insights
            )
            filename = f"{safe_name}_report.docx"
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"code": "INVALID_FORMAT", "message": "Format must be 'pdf' or 'docx'"},
            )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"Report generation failed for {dataset_id}: {exc}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"code": "REPORT_FAILED", "message": f"Report generation error: {exc}"},
        )

    logger.info(f"Report generated for {dataset_id}: {filename} ({len(content):,} bytes)")
    return content, filename
